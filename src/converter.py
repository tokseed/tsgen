"""
TypeScript Code Generator - Core Module
Поддержка GigaChat (Сбер) и OpenRouter API
"""

import os
import csv
import json
import base64
import uuid
import requests
import time
import hashlib
from typing import Optional, Dict, Any, List
from pathlib import Path
from datetime import datetime, timedelta

try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Импорт кэша
try:
    from cache import get_cache_manager

    CACHE_ENABLED = True
except ImportError:
    CACHE_ENABLED = False

    def get_cache_manager():
        return None


class GigaChatAuth:
    """
    GigaChat API авторизация (Сбер).

    Формат credentials: ClientID:ClientSecret
    Получите ключи в СберБизнес: https://developers.sber.ru/
    """

    AUTH_URL = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
    BASE_URL = "https://gigachat.devices.sberbank.ru/api/v1"
    MODELS_URL = "https://gigachat.devices.sberbank.ru/api/v1/models"

    def __init__(self, credentials: str, scope: str = "GIGACHAT_API_PERS"):
        self.credentials = credentials.strip()
        self.scope = scope
        self.access_token = None
        self.token_expires_at = 0

        # Проверяем формат: ClientID:ClientSecret (содержит двоеточие)
        self.is_valid_format = ":" in self.credentials

    def get_token(self) -> str:
        """Получение access token через OAuth."""
        # Кэширование токена
        if self.access_token and time.time() < self.token_expires_at - 60:
            return self.access_token

        if not self.is_valid_format:
            raise ValueError(
                "Неверный формат GIGACHAT_CREDENTIALS. "
                "Ожидается ClientID:ClientSecret (получите в СберБизнес)"
            )

        try:
            # Формируем Basic Auth: Base64(ClientID:ClientSecret)
            credentials_b64 = base64.b64encode(self.credentials.encode()).decode()

            headers = {
                "Content-Type": "application/x-www-form-urlencoded",
                "Accept": "application/json",
                "RqUID": str(uuid.uuid4()),
                "Authorization": f"Basic {credentials_b64}",
            }

            response = requests.post(
                self.AUTH_URL,
                headers=headers,
                data={"scope": self.scope},
                verify=False,
                timeout=30,
            )
            response.raise_for_status()

            result = response.json()
            self.access_token = result["access_token"]

            # Парсим время истечения (timestamp в секундах)
            expires_at = result.get("expires_at", 0)
            if isinstance(expires_at, (int, float)):
                self.token_expires_at = expires_at
            else:
                self.token_expires_at = time.time() + 1800

            return self.access_token

        except requests.exceptions.HTTPError as e:
            error_msg = f"GigaChat OAuth error: {e}"
            if hasattr(e, "response") and e.response is not None:
                error_msg += f" - {e.response.text}"
            raise ValueError(error_msg)
        except Exception as e:
            raise ValueError(f"GigaChat auth failed: {e}")

    def chat(self, messages: list, model: str = "GigaChat-2") -> str:
        """Отправка запроса к GigaChat."""
        token = self.get_token()

        response = requests.post(
            f"{self.BASE_URL}/chat/completions",
            headers={
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
            },
            json={
                "model": model,
                "messages": messages,
                "temperature": 0.3,
                "max_tokens": 4000,
            },
            verify=False,
            timeout=120,
        )
        response.raise_for_status()

        result = response.json()
        if "choices" not in result or len(result["choices"]) == 0:
            raise ValueError("Пустой ответ от GigaChat")

        return result["choices"][0]["message"]["content"]

    def check_connection(self) -> dict:
        """Проверка подключения к API."""
        try:
            token = self.get_token()
            response = requests.get(
                self.MODELS_URL,
                headers={"Authorization": f"Bearer {token}"},
                verify=False,
                timeout=10,
            )
            return {"status": "ok", "models": response.json()}
        except Exception as e:
            return {"status": "error", "error": str(e)}


class OpenRouterAuth:
    """OpenRouter API клиент с интеграцией LangFuse."""

    BASE_URL = "https://openrouter.ai/api/v1"

    MODELS = {
        "best": "anthropic/claude-3.5-sonnet",
        "balanced": "meta-llama/llama-3.3-70b-instruct",
        "fast": "meta-llama/llama-3.2-3b-instruct:free",
        "free": "google/gemma-2-9b-it:free",
        # Бесплатные модели OpenRouter
        "deepseek": "deepseek/deepseek-chat-v3-0324",
    }

    def __init__(self, api_key: str, model_preset: str = "balanced"):
        self.api_key = api_key
        self.model = self.MODELS.get(model_preset, self.MODELS["balanced"])

        # LangFuse интеграция
        self.langfuse = None
        self.trace = None
        self._init_langfuse()

    def _init_langfuse(self):
        """Инициализация LangFuse для трекинга."""
        try:
            from langfuse import Langfuse
            import os

            langfuse_public_key = os.getenv("LANGFUSE_PUBLIC_KEY")
            langfuse_secret_key = os.getenv("LANGFUSE_SECRET_KEY")
            langfuse_host = os.getenv("LANGFUSE_HOST", "https://cloud.langfuse.com")

            if langfuse_public_key and langfuse_secret_key:
                self.langfuse = Langfuse(
                    public_key=langfuse_public_key,
                    secret_key=langfuse_secret_key,
                    host=langfuse_host,
                )
                print("[LangFuse] Инициализирован успешно")
            else:
                print("[LangFuse] Ключи не найдены в .env")
        except ImportError:
            print("[LangFuse] Не установлен: pip install langfuse")
        except Exception as e:
            print(f"[LangFuse] Ошибка инициализации: {e}")

    def chat(self, messages: list, model: str = None) -> str:
        """Отправка запроса к OpenRouter с трекингом LangFuse."""
        model_to_use = model or self.model
        start_time = time.time()

        # Создаём trace в LangFuse
        if self.langfuse:
            try:
                self.trace = self.langfuse.trace(
                    name="openrouter-chat",
                    input={"messages": messages},
                    metadata={
                        "provider": "openrouter",
                        "model": model_to_use,
                    },
                )
            except Exception as e:
                print(f"[LangFuse] Ошибка создания trace: {e}")
                self.trace = None

        try:
            result = self._make_request(messages, model_to_use)

            # Логируем результат в LangFuse
            if self.langfuse and self.trace:
                try:
                    elapsed_time = time.time() - start_time
                    self.trace.update(
                        output={"content": result[:500] if result else None},
                        metadata={
                            "model": model_to_use,
                            "duration_ms": round(elapsed_time * 1000),
                        },
                    )
                    self.langfuse.flush()
                except Exception as e:
                    print(f"[LangFuse] Ошибка обновления trace: {e}")

            return result

        except Exception as e:
            # Логируем ошибку
            if self.langfuse and self.trace:
                try:
                    self.trace.update(
                        output={"error": str(e)},
                        level="ERROR",
                    )
                    self.langfuse.flush()
                except:
                    pass
            raise

    def _make_request(self, messages: list, model: str) -> str:
        """Внутренний метод для HTTP запроса к OpenRouter."""
        max_retries = 3
        retry_delay = 2

        for attempt in range(max_retries):
            try:
                response = requests.post(
                    f"{self.BASE_URL}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": os.getenv(
                            "APP_URL", "https://github.com/tsgen"
                        ),
                        "X-Title": "TS Generator",
                    },
                    json={
                        "model": model,
                        "messages": messages,
                        "temperature": 0.3,
                        "max_tokens": 4000,
                    },
                    timeout=120,
                )

                # Обработка ошибок API
                if response.status_code == 429:
                    raise ValueError("Rate limit превышен. Подождите немного.")
                elif response.status_code == 401:
                    raise ValueError("Неверный API ключ OpenRouter.")
                elif response.status_code == 402:
                    raise ValueError("Недостаточно кредитов на счету OpenRouter.")

                response.raise_for_status()
                result = response.json()

                if "error" in result:
                    error_msg = result["error"].get("message", "Неизвестная ошибка API")
                    raise ValueError(f"Ошибка OpenRouter API: {error_msg}")

                if "choices" not in result or len(result["choices"]) == 0:
                    raise ValueError("Пустой ответ от OpenRouter")

                content = result["choices"][0]["message"]["content"]

                if not content or content.strip() == "":
                    raise ValueError("Пустое содержимое ответа")

                # Логируем usage в консоль для отладки
                usage = result.get("usage", {})
                if usage:
                    print(
                        f"[OpenRouter] Tokens: {usage.get('prompt_tokens', 0)} in / {usage.get('completion_tokens', 0)} out"
                    )

                return content

            except requests.exceptions.Timeout:
                if attempt == max_retries - 1:
                    raise ValueError("Таймаут соединения с OpenRouter.")
                time.sleep(retry_delay * (attempt + 1))
            except requests.exceptions.RequestException as e:
                if attempt == max_retries - 1:
                    raise ValueError(f"Ошибка соединения: {str(e)}")
                time.sleep(retry_delay)

        raise ValueError("Не удалось выполнить запрос после нескольких попыток")

    def check_connection(self) -> dict:
        """Проверка подключения к API."""
        try:
            response = requests.get(
                f"{self.BASE_URL}/models",
                headers={"Authorization": f"Bearer {self.api_key}"},
                timeout=10,
            )
            if response.status_code == 200:
                return {"status": "ok", "models": response.json()}
            return {"status": "error", "error": f"HTTP {response.status_code}"}
        except Exception as e:
            return {"status": "error", "error": str(e)}


class FileProcessor:
    """Обработчик файлов."""

    SUPPORTED_FORMATS = {
        ".csv": "csv",
        ".xls": "excel",
        ".xlsx": "excel",
        ".pdf": "pdf",
        ".docx": "docx",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
    }

    @classmethod
    def get_file_type(cls, filepath: str) -> Optional[str]:
        """Определение типа файла по расширению."""
        return cls.SUPPORTED_FORMATS.get(Path(filepath).suffix.lower())

    @classmethod
    def process_file(cls, filepath: str) -> Dict[str, Any]:
        """Обработка файла."""
        file_type = cls.get_file_type(filepath)
        if not file_type:
            raise ValueError(f"Неподдерживаемый формат: {Path(filepath).suffix}")

        processors = {
            "csv": cls._process_csv,
            "excel": cls._process_excel,
            "pdf": cls._process_pdf,
            "docx": cls._process_docx,
            "image": cls._process_image,
        }

        processor = processors.get(file_type)
        if not processor:
            raise ValueError(f"Нет обработчика для: {file_type}")

        content, structure = processor(filepath)
        return {"type": file_type, "content": content, "structure": structure}

    @staticmethod
    def _process_csv(filepath: str) -> tuple:
        """Обработка CSV через встроенный модуль csv."""
        with open(filepath, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return "", {"columns": [], "rows_count": 0, "dtypes": {}}

        headers = rows[0]
        data_rows = rows[1:50]  # Первые 50 строк данных

        # Определяем типы по первым значениям
        dtypes = {}
        if data_rows:
            for i, header in enumerate(headers):
                sample = data_rows[0][i] if i < len(data_rows[0]) else ""
                if sample.isdigit():
                    dtypes[header] = "int64"
                elif sample.replace(".", "").replace("-", "").isdigit():
                    dtypes[header] = "float64"
                elif sample.lower() in ("true", "false"):
                    dtypes[header] = "bool"
                else:
                    dtypes[header] = "object"

        content_lines = [",".join(headers)]
        content_lines.extend([",".join(row) for row in data_rows])

        structure = {
            "columns": headers,
            "rows_count": len(data_rows),
            "dtypes": dtypes,
        }
        return "\n".join(content_lines), structure

    @staticmethod
    def _process_excel(filepath: str) -> tuple:
        """Обработка Excel через openpyxl."""
        try:
            from openpyxl import load_workbook

            wb = load_workbook(filepath, read_only=True, data_only=True)
            ws = wb.active

            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 51:  # Заголовок + 50 строк
                    break
                rows.append([str(cell) if cell is not None else "" for cell in row])
            wb.close()

            if not rows:
                return "", {"columns": [], "rows_count": 0, "dtypes": {}}

            headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
            data_rows = rows[1:]

            dtypes = {}
            if data_rows:
                for i, header in enumerate(headers):
                    sample = data_rows[0][i] if i < len(data_rows[0]) else ""
                    if sample.isdigit():
                        dtypes[header] = "int64"
                    elif sample.replace(".", "").replace("-", "").isdigit():
                        dtypes[header] = "float64"
                    else:
                        dtypes[header] = "object"

            content_lines = [",".join(headers)]
            content_lines.extend([",".join(row) for row in data_rows])

            structure = {
                "columns": headers,
                "rows_count": len(data_rows),
                "dtypes": dtypes,
            }
            return "\n".join(content_lines), structure
        except ImportError:
            return "Excel processing requires openpyxl", {
                "error": "openpyxl not installed"
            }
        except Exception as e:
            return f"Excel error: {e}", {"error": str(e)}

    @staticmethod
    def _process_pdf(filepath: str) -> tuple:
        reader = PdfReader(filepath)
        text = "".join(page.extract_text() or "" for page in reader.pages[:10])
        structure = {"pages": len(reader.pages), "text_length": len(text)}
        return text[:5000], structure

    @staticmethod
    def _process_docx(filepath: str) -> tuple:
        """Обработка DOCX через python-docx."""
        if not DOCX_AVAILABLE:
            return "DOCX processing requires python-docx", {
                "error": "python-docx not installed"
            }

        try:
            doc = Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs[:100]])

            tables_data = []
            for table in doc.tables[:5]:
                tables_data.append(
                    [[cell.text for cell in row.cells] for row in table.rows]
                )

            structure = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "text_length": len(text),
            }

            content = text
            if tables_data:
                content += "\n\n=== TABLES ===\n"
                for i, table in enumerate(tables_data):
                    content += f"\nTable {i + 1}:\n"
                    for row in table:
                        content += " | ".join(row) + "\n"

            return content[:5000], structure
        except Exception as e:
            return f"DOCX error: {e}", {"error": str(e)}

    @staticmethod
    def _process_image(filepath: str) -> tuple:
        """Обработка изображений через Pillow (OCR требует tesseract)."""
        if not PIL_AVAILABLE:
            return "Image processing requires Pillow", {"error": "Pillow not installed"}

        try:
            image = Image.open(filepath)
            structure = {"size": image.size, "mode": image.mode, "format": image.format}

            # Пытаемся использовать OCR если доступен
            try:
                import pytesseract

                text = pytesseract.image_to_string(image, lang="rus+eng")
                return text[:5000], {**structure, "text_length": len(text)}
            except (ImportError, Exception):
                # Если OCR недоступен, возвращаем метаданные
                return (
                    f"[Image: {image.size[0]}x{image.size[1]}, {image.format}]",
                    structure,
                )
        except Exception as e:
            return f"Image error: {e}", {"error": str(e)}


class TypeScriptCodeGenerator:
    """Генератор TypeScript кода."""

    SYSTEM_PROMPT = """Ты — старший TypeScript разработчик в Сбер. Твоя задача — создать профессиональный TypeScript код для парсинга файлов и преобразования данных в JSON.

## ВХОДНЫЕ ДАННЫЕ:
- Тип файла: {file_type}
- Структура данных: {structure}
- Пример содержимого (первые строки/символы):
{content}

## ЦЕЛЕВАЯ JSON СХЕМА (обязательно для соблюдения):
{target_json}

## ТЕХНИЧЕСКОЕ ЗАДАНИЕ (СТРОГОЕ ТРЕБОВАНИЕ):

### 1. АРХИТЕКТУРА КОДА

#### 1.1 Интерфейсы
- Создай ПОЛНЫЕ TypeScript интерфейсы для ВСЕХ структур данных
- Используй ТОЧНЫЕ типы: string, number, boolean, null | undefined
- КАТЕГОРИЧЕСКИ избегай `any` — используй `unknown` или конкретные типы
- Для необязательных полей используй `?` (optional modifier)
- Для nullable полей используй `| null`
- Экспортируй ВСЕ публичные интерфейсы: `export interface`

#### 1.2 Основная функция
- Сигнатура (ОБЯЗАТЕЛЬНА): `export async function transformData(input: string): Promise<{interface_name}>`
- Функция должна быть `async` даже если не использует await (для единообразия)
- Возвращаемый тип должен ТОЧНО соответствовать целевой JSON схеме

### 2. ПАРСИНГ И ПРЕОБРАЗОВАНИЕ

#### 2.1 Обработка входных данных
- Определи формат входных данных по типу файла ({file_type})
- Для CSV/Excel: раздели на строки, затем на колонки
- Для JSON: распарси через JSON.parse()
- Для текста: извлеки структурированные данные

#### 2.2 Преобразование типов
- Строки в числа: `Number(value)` или `parseFloat(value)` или `parseInt(value, 10)`
- Строки в булевы: `value.toLowerCase() === 'true'` или `value === '1'`
- Пустые значения: `value.trim() === '' ? null : value`
- Обработка NaN: `const num = Number(val); if (isNaN(num)) throw new Error(...)`

#### 2.3 Валидация данных
- Проверь ВСЕ обязательные поля из целевой схемы
- Проверь типы данных после преобразования
- Выбрасывай ошибку при несоответствии схемы

### 3. ОБРАБОТКА ОШИБОК (ОБЯЗАТЕЛЬНО)

#### 3.1 Try-catch блоки
- Оберни КАЖДУЮ операцию парсинга в try-catch
- Логируй ошибки с контекстом (номер строки, имя поля)

#### 3.2 Типы ошибок
- Пустой файл: `throw new Error('Пустой файл: входные данные не содержат строк')`
- Неверный формат: `throw new Error('Неверный формат: ожидался CSV, получен...')`
- Missing поля: `throw new Error('Отсутствует обязательное поле: fieldName')`
- Ошибка типа: `throw new Error('Неверный тип поля fieldName: ожидался number, получен string')`

### 4. КАЧЕСТВО КОДА

#### 4.1 Стиль кода
- БЕЗ внешних библиотек (только стандартная библиотека ES2020+)
- Используй const/let, избегай var
- Стрелочные функции для callback'ов
- Template literals для интерполяции строк

#### 4.2 ВАЖНО: БЕЗ КОММЕНТАРИЕВ!
- НИКАКИХ комментариев в коде (// или /* */)
- Код должен быть самодокументируемым
- Понятные имена переменных вместо комментариев

#### 4.3 Именование
- camelCase для переменных и функций
- PascalCase для интерфейсов и типов
- Понятные имена: `userData`, `transformData`, `parseCSV`

### 5. ФОРМАТ ВЫВОДА (КРИТИЧНО)

- Верни ТОЛЬКО TypeScript код в markdown блоке
- Формат: ```typescript ... ```
- БЕЗ объяснений, БЕЗ текста до или после кода
- БЕЗ комментариев в коде
- Код должен быть ГОТОВ К ПРОДАКШЕНУ

### 6. ПРИМЕР СТРУКТУРЫ ОТВЕТА:

```typescript
// Интерфейсы для типизации
export interface TargetData {{
  // поля из JSON схемы
}}

export interface Metadata {{
  sourceType: string;
  rowCount: number;
  processedAt: string;
}}

/**
 * Парсит входные данные и преобразует в формат JSON
 * @param input - содержимое файла в виде строки
 * @returns Promise<TargetData> с преобразованными данными
 * @throws Error при неверном формате или отсутствии обязательных полей
 */
export async function transformData(input: string): Promise<TargetData> {{
  try {{
    // 1. Проверка пустого ввода
    if (!input || input.trim() === '') {{
      throw new Error('Пустой файл: входные данные не содержат строк');
    }}

    // 2. Парсинг входных данных
    const lines = input.trim().split('\\n');
    const headers = lines[0].split(',').map(h => h.trim());
    
    // 3. Преобразование в целевую схему
    const result: TargetData = {{
      data: [],
      metadata: {{
        sourceType: '{file_type}',
        rowCount: lines.length - 1,
        processedAt: new Date().toISOString(),
      }},
    }};
    
    // 4. Обработка каждой строки
    for (let i = 1; i < lines.length; i++) {{
      const values = lines[i].split(',').map(v => v.trim());
      // ... преобразование
    }}
    
    return result;
  }} catch (error) {{
    if (error instanceof Error) {{
      throw new Error(`Ошибка трансформации: ${{error.message}}`);
    }}
    throw new Error('Неизвестная ошибка при трансформации');
  }}
}}

// Вспомогательные функции
function parseNumber(value: string, fieldName: string): number {{
  const num = Number(value);
  if (isNaN(num)) {{
    throw new Error(`Неверный тип поля ${{fieldName}}: ожидался number, получен "${{value}}"`);
  }}
  return num;
}}
```
"""

    def __init__(self, llm_config: Optional[Dict[str, Any]] = None):
        self.llm_config = llm_config or {}
        self.provider = self.llm_config.get("provider", "auto")
        self.gigachat = None
        self.openrouter = None
        self.cache_manager = get_cache_manager() if CACHE_ENABLED else None
        self._init_llm()

    def _init_llm(self):
        """Инициализация LLM провайдера."""
        if self.provider == "gigachat" or self.provider == "auto":
            credentials = self.llm_config.get("credentials", "")
            if credentials:
                try:
                    self.gigachat = GigaChatAuth(credentials)
                    self.provider = "gigachat"
                    return
                except Exception as e:
                    print(f"Warning: Failed to init GigaChat: {e}")

        if self.provider == "openrouter" or self.provider == "auto":
            api_key = self.llm_config.get("api_key", "")
            if api_key:
                try:
                    self.openrouter = OpenRouterAuth(api_key)
                    self.provider = "openrouter"
                    return
                except Exception as e:
                    print(f"Warning: Failed to init OpenRouter: {e}")

        print("Warning: No LLM credentials, using mock mode")
        self.provider = "mock"

    def _generate_cache_key(self, file_data: Dict[str, Any], target_json: str) -> str:
        """Генерация ключа кэша на основе входных данных."""
        cache_data = {
            "file_type": file_data.get("type", "unknown"),
            "structure": file_data.get("structure", {}),
            "content_preview": file_data.get("content", "")[:1000],
            "target_json": target_json,
        }
        cache_string = json.dumps(cache_data, sort_keys=True, ensure_ascii=False)
        return hashlib.sha256(cache_string.encode()).hexdigest()

    def generate(self, file_data: Dict[str, Any], target_json: str) -> str:
        """Генерация TypeScript кода."""
        # Проверка кэша
        if self.cache_manager:
            cache_key = self._generate_cache_key(file_data, target_json)
            cached_result = self.cache_manager.get(cache_key)
            if cached_result:
                print(f"[CACHE HIT] Возвращаем результат из кэша")
                return cached_result
            print(f"[CACHE MISS] Генерируем новый результат")

        if self.provider == "mock" or (not self.gigachat and not self.openrouter):
            result = self._mock_generate(file_data, target_json)
            # Сохранение в кэш даже mock результатов
            if self.cache_manager:
                cache_key = self._generate_cache_key(file_data, target_json)
                self.cache_manager.set(cache_key, result)
            return result

        try:
            target_dict = json.loads(target_json)
            interface_name = self._infer_interface_name(target_dict)
        except json.JSONDecodeError:
            interface_name = "TargetData"

        messages = [
            {
                "role": "system",
                "content": "Ты экспертный TypeScript разработчик. Верни ТОЛЬКО код без объяснений.",
            },
            {
                "role": "user",
                "content": self.SYSTEM_PROMPT.format(
                    file_type=file_data["type"],
                    structure=json.dumps(file_data["structure"], ensure_ascii=False),
                    content=file_data["content"][:3000],
                    target_json=target_json,
                    interface_name=interface_name,
                ),
            },
        ]

        # Используем доступный LLM
        if self.openrouter:
            result = self.openrouter.chat(messages)
        elif self.gigachat:
            result = self.gigachat.chat(messages)
        else:
            result = self._mock_generate(file_data, target_json)

        result = self._clean_code(result)

        # Сохранение в кэш
        if self.cache_manager:
            cache_key = self._generate_cache_key(file_data, target_json)
            self.cache_manager.set(cache_key, result)

        return result

    def _mock_generate(self, file_data: Dict[str, Any], target_json: str) -> str:
        """Генерация mock кода."""
        try:
            target_dict = json.loads(target_json)
            interface_name = self._infer_interface_name(target_dict)
            interface_def = self._generate_interface(target_dict, interface_name)
        except json.JSONDecodeError:
            interface_name = "TargetData"
            interface_def = "interface TargetData {{ [key: string]: any; }}"

        file_type = file_data.get("type", "unknown")

        return f"""// TypeScript Code (Mock Mode)
{interface_def}

async function transformData(input: string): Promise<{interface_name}> {{
  try {{
    const lines = input.trim().split('\\n');
    if (lines.length === 0) throw new Error('Пустой файл');

    const result: {interface_name} = {{
      data: [],
      metadata: {{
        sourceType: '{file_type}',
        columns: lines[0]?.split(',').map(c => c.trim()) || [],
        rowCount: Math.max(0, lines.length - 1),
        processedAt: new Date().toISOString(),
      }},
    }};

    const headers = result.metadata.columns;
    for (let i = 1; i < lines.length; i++) {{
      const values = lines[i].split(',').map(v => v.trim().replace(/^"|"$/g, ''));
      if (values.length !== headers.length) continue;

      const row: {{ [key: string]: string }} = {{}};
      headers.forEach((header, idx) => {{ row[header] = values[idx] || ''; }});
      result.data.push(row as any);
    }}

    return result;
  }} catch (error) {{
    throw new Error(`Ошибка: ${{error instanceof Error ? error.message : String(error)}}`);
  }}
}}

export {{ transformData }};
"""

    @staticmethod
    def _infer_interface_name(data: Dict) -> str:
        """Определение имени интерфейса."""
        return "TransformResult" if "data" in data else "TargetData"

    @staticmethod
    def _generate_interface(data: Dict, name: str, indent: int = 0) -> str:
        """Генерация TypeScript интерфейса."""
        # Если data — список, берём первый элемент
        if isinstance(data, list):
            if not data:
                return f"interface {name} {{ [key: string]: any; }}"
            data = (
                data[0]
                if isinstance(data[0], dict)
                else {"value": type(data[0]).__name__}
            )

        prefix = "  " * indent
        lines = [f"{prefix}interface {name} {{"]

        for key, value in data.items():
            if isinstance(value, dict):
                nested_name = f"{name}{key.capitalize()}"
                lines.append(
                    TypeScriptCodeGenerator._generate_interface(
                        value, nested_name, indent + 1
                    )
                )
                lines.append(f"{prefix}  {key}: {nested_name};")
            elif isinstance(value, list):
                if value and isinstance(value[0], dict):
                    nested_name = f"{name}{key.capitalize()}Item"
                    lines.append(
                        TypeScriptCodeGenerator._generate_interface(
                            value[0], nested_name, indent + 1
                        )
                    )
                    lines.append(f"{prefix}  {key}: {nested_name}[];")
                else:
                    item_type = "any"
                    if value:
                        item_type = type(value[0]).__name__
                        item_type = {
                            "str": "string",
                            "int": "number",
                            "float": "number",
                            "bool": "boolean",
                        }.get(item_type, "any")
                    lines.append(f"{prefix}  {key}: {item_type}[];")
            else:
                ts_type = type(value).__name__
                ts_type = {
                    "str": "string",
                    "int": "number",
                    "float": "number",
                    "bool": "boolean",
                }.get(ts_type, "any")
                lines.append(f"{prefix}  {key}: {ts_type};")

        lines.append(f"{prefix}}}")
        return "\n".join(lines)

    @staticmethod
    def _clean_code(code: str) -> str:
        """Очистка кода от markdown."""
        code = code.strip()
        for prefix in ["```typescript", "```ts", "```"]:
            if code.startswith(prefix):
                code = code[len(prefix) :]
                break
        if code.endswith("```"):
            code = code[:-3]
        return code.strip()


class ConverterAgent:
    """Конвертер файлов в TypeScript."""

    def __init__(self, llm_config: Optional[Dict[str, Any]] = None):
        self.file_processor = FileProcessor()
        self.code_generator = TypeScriptCodeGenerator(llm_config)

    def process(self, filepath: str, target_json: str) -> Dict[str, Any]:
        """Обработка файла."""
        file_data = self.file_processor.process_file(filepath)
        typescript_code = self.code_generator.generate(file_data, target_json)

        return {
            "typescript_code": typescript_code,
            "file_info": {
                "type": file_data["type"],
                "structure": file_data["structure"],
            },
        }


def generate_typescript(
    filepath: str, target_json: str, llm_config: Optional[Dict] = None
) -> str:
    """Генерация TypeScript кода."""
    agent = ConverterAgent(llm_config)
    return agent.process(filepath, target_json)["typescript_code"]
